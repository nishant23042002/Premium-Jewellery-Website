from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


OUT = Path("docs/Ambika_Jewellers_Client_Deployment_Report.docx")
PDF_OUT = Path("output/pdf/Ambika_Jewellers_Client_Deployment_Report.pdf")
REPORT_DATE = date.today().strftime("%d %B %Y")

NAVY = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY = RGBColor(85, 85, 85)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
WARN = "FFF3CD"
NOTE = "EEF7FF"
BEST = "EAF6EF"
RISK = "FDECEC"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D9E2EC"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(col)


def run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_para(doc, text="", style=None, bold=False, italic=False, color=None, size=None, align=None, before=None, after=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if before is not None:
        p.paragraph_format.space_before = Pt(before)
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    if text:
        r = p.add_run(text)
        run_font(r, size=size, color=color, bold=bold, italic=italic)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        run_font(r, size={1: 16, 2: 13, 3: 12}.get(level, 11), color=BLUE if level < 3 else DARK_BLUE, bold=True)
    return p


def add_callout(doc, label, text, fill):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [6.5])
    set_table_borders(table, "DADCE0")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"{label}: ")
    run_font(r, size=10.5, bold=True, color=NAVY)
    r = p.add_run(text)
    run_font(r, size=10.5)
    doc.add_paragraph()


def add_matrix(doc, headers, rows, widths=None, header_fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=len(headers))
    if widths is None:
        widths = [6.5 / len(headers)] * len(headers)
    set_table_width(table, widths)
    set_table_borders(table)
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_shading(hdr[i], header_fill)
        p = hdr[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        run_font(r, bold=True, size=9.5, color=NAVY)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(val))
            run_font(r, size=9.2)
    return table


PHASES = [
    ("Phase 1", "Business Preparation", ["business email", "Google Workspace", "account ownership", "password manager", "2FA", "recovery planning", "login inventory"], "Client"),
    ("Phase 2", "Domain Management", ["domain choice", "registrar", "TLD", "DNS", "nameservers", "TTL", "renewal", "privacy", "transfer lock"], "Client"),
    ("Phase 3", "Hosting", ["Vercel", "plans", "domain connection", "environment variables", "preview deployments", "rollback", "logs", "monitoring"], "Developer"),
    ("Phase 4", "Database", ["MongoDB Atlas", "organization", "project", "cluster", "users", "roles", "IP allowlist", "backups", "alerts"], "Developer"),
    ("Phase 5", "Cloudinary", ["media folders", "API keys", "upload presets", "optimization", "transformations", "naming", "backup"], "Developer"),
    ("Phase 6", "Firebase", ["project", "authentication", "Google login", "Phone OTP", "authorized domains", "security rules", "environment variables"], "Developer"),
    ("Phase 7", "Payment Gateway", ["Razorpay", "KYC", "PAN", "GST", "current account", "settlements", "webhooks", "live keys", "refunds"], "Client"),
    ("Phase 8", "Google Services", ["Search Console", "Analytics", "Tag Manager", "Business Profile", "Maps", "Merchant Center", "PageSpeed"], "Developer"),
    ("Phase 9", "SEO Setup", ["metadata", "Open Graph", "schema", "local business", "product schema", "image SEO", "canonical", "sitemap", "robots"], "Developer"),
    ("Phase 10", "Deployment Checklist", ["environment", "database", "Cloudinary", "Firebase", "payments", "SSL", "domain", "testing", "accessibility"], "Developer"),
    ("Phase 11", "Testing", ["desktop", "tablet", "mobile", "Android", "iPhone", "payment", "OTP", "forms", "links", "errors", "speed"], "Developer"),
    ("Phase 12", "Security", ["HTTPS", "headers", "rate limiting", "secrets", "XSS", "CSRF", "validation", "uploads", "logging", "backups"], "Developer"),
    ("Phase 13", "Maintenance", ["weekly", "monthly", "quarterly", "database backups", "updates", "analytics", "broken links", "performance"], "Shared"),
    ("Phase 14", "Client Handover", ["source code", "GitHub", "deployment access", "credentials", "documentation", "PDF guide", "walkthrough", "support"], "Developer"),
    ("Phase 15", "Future Scaling", ["admin dashboard", "inventory", "orders", "CRM", "WhatsApp", "SMS", "email marketing", "reviews", "loyalty"], "Shared"),
    ("Phase 16", "Estimated Annual Running Cost", ["domain", "hosting", "Cloudinary", "MongoDB", "Firebase", "Workspace", "Razorpay", "maintenance"], "Client"),
    ("Phase 17", "Client Account Ownership Matrix", ["service", "purpose", "owner", "developer access", "client access", "recovery email", "criticality"], "Shared"),
    ("Phase 18", "Complete Deployment Flowchart", ["business decision", "accounts", "domain", "hosting", "database", "media", "auth", "payments", "analytics", "launch"], "Developer"),
    ("Phase 19", "Developer Deployment SOP", ["pre-deployment", "deployment", "launch day", "post-launch", "rollback", "bugs", "maintenance"], "Developer"),
    ("Phase 20", "Appendix", ["naming", "folders", "Git workflow", "branches", "environment variables", "password policy", "tools", "templates"], "Shared"),
]


def configure_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
    for name in ["List Bullet", "List Number"]:
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(11)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.text = "Ambika Jewellers - Client Deployment Handbook"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for r in header.runs:
        run_font(r, size=8.5, color=GRAY)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.text = "Confidential handover guide | Prepared for jewellery website launch"
    for r in footer.runs:
        run_font(r, size=8.5, color=GRAY)


def cover(doc):
    add_para(doc, "Premium Jewellery Website", bold=True, color=GRAY, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, after=14)
    add_para(doc, "Client Deployment Report", bold=True, color=NAVY, size=30, align=WD_ALIGN_PARAGRAPH.CENTER, after=5)
    add_para(doc, "A complete agency handbook for deployment, launch, handover, and maintenance", color=GRAY, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, after=34)
    add_callout(doc, "Scope", "This guide assumes a Next.js, TypeScript, Tailwind CSS, Node.js, MongoDB Atlas, Cloudinary, Firebase Authentication, Razorpay, and Vercel stack. It is written for a non-technical jewellery business owner and a delivery team that needs a repeatable deployment SOP.", NOTE)
    add_matrix(
        doc,
        ["Prepared For", "Project Type", "Prepared Date", "Primary Outcome"],
        [["Ambika Jewellers", "Premium business website", REPORT_DATE, "Website live with client-owned accounts, documented access, launch QA, and maintenance plan"]],
        [1.55, 1.55, 1.2, 2.2],
    )
    add_para(doc, "How to Use This Report", style="Heading 1")
    add_bullets(doc, [
        "Client should read Phases 1, 2, 7, 13, 14, 16, and 17 before purchasing or approving any service.",
        "Developer should follow Phases 3 through 12 in sequence and record every credential, DNS change, key, webhook, and launch decision.",
        "Both parties should sign off on the launch checklist, account ownership matrix, and support arrangement before public launch.",
        "Do not launch with temporary personal accounts, shared OTP numbers, missing recovery options, unpaid renewals, or unverified payment settlement details.",
    ])
    doc.add_page_break()


def toc(doc):
    add_heading(doc, "Table of Contents", 1)
    add_bullets(doc, [f"{num}: {title}" for num, title, _, _ in PHASES])
    add_bullets(doc, [
        "Deployment Timeline",
        "Risk Matrix",
        "Maintenance Calendar",
        "Final Summary",
    ])
    add_callout(doc, "Note", "Page numbers may update when opened in Microsoft Word. The structure is intentionally phase-based so the team can use headings, search, and navigation pane items during the handover meeting.", NOTE)
    doc.add_page_break()


def phase_content(doc, num, title, topics, owner):
    add_heading(doc, f"{num}: {title}", 1)
    add_callout(doc, "Objective", f"Complete the {title.lower()} work so the jewellery website can move forward without account ownership gaps, launch blockers, or maintenance ambiguity.", NOTE)
    add_matrix(
        doc,
        ["Owner", "Estimated Cost", "Estimated Time", "Prerequisites"],
        [[owner, cost_for(title), time_for(title), prerequisites_for(title)]],
        [1.1, 1.4, 1.2, 2.8],
    )

    add_heading(doc, "Why This Phase Is Required", 2)
    add_para(doc, detailed_why(title, topics))
    add_para(doc, "For a premium jewellery store, trust is part of the brand. Domain ownership, authentication, payment verification, media quality, page speed, and support readiness all influence whether a visitor feels confident enough to enquire, book a visit, or place an order.")

    add_heading(doc, "Responsibilities", 2)
    add_matrix(
        doc,
        ["Activity", "Client Responsibility", "Developer Responsibility"],
        [[topic.title(), client_resp(topic, owner), dev_resp(topic)] for topic in topics[:7]],
        [1.8, 2.35, 2.35],
    )

    add_heading(doc, "Step-by-Step Procedure", 2)
    add_numbers(doc, procedure_steps(title, topics))

    add_heading(doc, "Detailed Implementation Notes", 2)
    for topic in topics:
        add_heading(doc, topic.title(), 3)
        add_para(doc, topic_detail(title, topic))
        add_bullets(doc, topic_actions(title, topic))

    add_heading(doc, "Common Mistakes", 2)
    add_bullets(doc, mistakes_for(title, topics))
    add_callout(doc, "Warning", warning_for(title), WARN)

    add_heading(doc, "Best Practices", 2)
    add_bullets(doc, best_for(title, topics))
    add_callout(doc, "Best Practice", best_callout(title), BEST)

    add_heading(doc, "Checklist Before Moving to Next Phase", 2)
    add_bullets(doc, checklist_for(title, topics))

    add_heading(doc, "Decision Tree", 2)
    add_para(doc, decision_tree(title), style=None)

    add_heading(doc, "Agency Notes", 2)
    add_para(doc, agency_notes(title, topics))
    doc.add_page_break()


def topic_detail(title, topic):
    return (
        f"For {topic}, the delivery team should confirm the business purpose, owner, access level, billing impact, "
        f"production value, and recovery path. In the context of {title.lower()}, this item must be explained to the "
        "client in plain language: what it controls, what can break if it is misconfigured, what renewal or usage cost "
        "may apply, and who can restore service during an emergency. Record the final setting, not only the intention."
    )


def topic_actions(title, topic):
    return [
        f"Create a documented record for {topic}: account URL, owner email, access role, support contact, and renewal or billing model.",
        f"Verify {topic} in a production-equivalent environment and capture the test result in the launch checklist.",
        f"Confirm whether {topic} contains secrets, payment data, customer data, or business-critical configuration.",
        f"Add {topic} to the monthly maintenance review if it can affect uptime, search visibility, orders, enquiries, or brand trust.",
    ]


def cost_for(title):
    mapping = {
        "Business Preparation": "INR 0-10,000/year",
        "Domain Management": "INR 800-3,500/year",
        "Hosting": "INR 0-20,000+/year",
        "Database": "INR 0-25,000+/year",
        "Cloudinary": "INR 0-20,000+/year",
        "Firebase": "INR 0-10,000+/year",
        "Payment Gateway": "Usually no setup fee; transaction fees apply",
        "Google Services": "Mostly free; ads optional",
        "SEO Setup": "Included in development or INR 10,000-50,000 setup",
        "Maintenance": "INR 3,000-25,000/month",
    }
    return mapping.get(title, "Depends on scope; document before approval")


def time_for(title):
    if title in ["Business Preparation", "Domain Management", "Payment Gateway"]:
        return "1-7 business days"
    if title in ["Testing", "Deployment Checklist", "Client Handover"]:
        return "1-3 business days"
    return "2-6 hours after prerequisites"


def prerequisites_for(title):
    base = {
        "Business Preparation": "Owner phone, recovery email, business identity, account naming decision",
        "Domain Management": "Final brand name, registered email, payment method",
        "Hosting": "GitHub repository, production env vars, verified domain decision",
        "Database": "MongoDB account, data model, security roles, backup policy",
        "Payment Gateway": "Legal business documents, bank account, KYC, GST/PAN details",
    }
    return base.get(title, "Previous phase sign-off, client-owned accounts, documented approvals")


def detailed_why(title, topics):
    return (
        f"{title} is required because it establishes the operational foundation for {', '.join(topics[:5])}. "
        "If this phase is rushed, the website may still appear live, but the business can later lose access, miss renewals, expose secrets, fail verification, or struggle to diagnose customer-facing issues."
    )


def client_resp(topic, owner):
    if owner == "Client" or topic in ["KYC", "PAN", "GST", "current account", "domain choice", "business email"]:
        return "Own account, approve details, pay fees, keep recovery access."
    if owner == "Shared":
        return "Approve business rules and retain final ownership."
    return "Provide credentials through secure handover and approve production settings."


def dev_resp(topic):
    return f"Configure {topic}, test it in production-like conditions, document access, and explain the result in plain language."


def procedure_steps(title, topics):
    return [
        f"Confirm the business decision and owner for {title.lower()} before creating or changing any account.",
        f"Create or open the required service and use the client-owned business email wherever possible.",
        f"Configure the core settings for {', '.join(topics[:4])}.",
        "Record account URL, login email, recovery email, billing owner, renewal date, and emergency contact.",
        "Apply security controls: strong unique password, 2FA, least-privilege developer access, and recovery backup.",
        "Connect the service to the website only after test credentials or staging settings are verified.",
        "Run a live or production-equivalent validation and capture proof in the launch checklist.",
        "Get client sign-off before proceeding to the next phase.",
    ]


def mistakes_for(title, topics):
    return [
        "Using the developer's personal email for a client-owned account.",
        "Sharing passwords through chat apps instead of a password manager or secure handover method.",
        f"Configuring {topics[0]} once and not documenting who can recover it later.",
        "Leaving trial settings, test keys, staging URLs, placeholder metadata, or default security settings in production.",
        "Skipping renewal, billing, alert, and emergency contact setup because the website appears to work today.",
    ]


def warning_for(title):
    return f"Do not treat {title.lower()} as a one-time setup task. Every account, key, domain record, payment setting, and recovery option must remain usable by the client after the developer exits."


def best_for(title, topics):
    return [
        "Use client-owned accounts with developer access granted separately.",
        "Store credentials in a password manager and record recovery options in the handover register.",
        "Use production, preview, and development environments separately; never mix live keys with local testing.",
        f"Name resources clearly, for example ambika-jewellers-production-{topics[0].replace(' ', '-')}.",
        "Enable alerts for billing, downtime, security, database usage, failed webhooks, and domain renewal.",
        "Schedule a 30-minute client walkthrough after each business-critical service is configured.",
    ]


def best_callout(title):
    return f"The safest agency practice is to make {title.lower()} client-owned, developer-assisted, and fully documented before the public launch date."


def checklist_for(title, topics):
    return [
        f"[ ] Owner confirmed for {title}.",
        "[ ] Billing and renewal responsibility documented.",
        "[ ] Recovery email and phone number tested.",
        "[ ] 2FA enabled and backup codes stored securely.",
        f"[ ] {', '.join(t.title() for t in topics[:3])} configured and verified.",
        "[ ] Production setting reviewed against staging or test setting.",
        "[ ] Screenshot or written proof added to launch records.",
        "[ ] Client sign-off received.",
    ]


def decision_tree(title):
    return (
        "Start -> Is the account client-owned? If no, create/migrate before launch. "
        "If yes -> Is 2FA enabled and recovery tested? If no, complete security setup. "
        f"If yes -> Are {title.lower()} settings documented and production-tested? If no, validate. "
        "If yes -> proceed to the next phase."
    )


def agency_notes(title, topics):
    return (
        f"During {title.lower()}, communicate in non-technical terms. Explain what each service does, what the client must pay for, "
        f"what happens if it expires, and who should be contacted during an emergency. For topics such as {', '.join(topics)}, "
        "avoid jargon until the client understands the business impact. The agency should leave behind a clean operating trail: "
        "account owner, access level, renewal date, support email, emergency action, and expected monthly maintenance check."
    )


def special_sections(doc):
    add_heading(doc, "Deployment Timeline", 1)
    add_matrix(doc, ["Stage", "Typical Duration", "Primary Owner", "Exit Criteria"], [
        ["Preparation and ownership", "1-3 days", "Client", "Business email, recovery, password manager, domain decision ready"],
        ["Infrastructure setup", "1 day", "Developer", "Vercel, MongoDB, Cloudinary, Firebase configured"],
        ["Payment and Google services", "2-7 days", "Shared", "Razorpay live-ready, Search Console and Analytics verified"],
        ["Launch QA", "1-2 days", "Developer", "Responsive, payment, auth, SEO, security, and performance checks passed"],
        ["Handover and support", "1 day", "Shared", "Access matrix, guide, walkthrough, and support window agreed"],
    ], [1.6, 1.2, 1.2, 2.5])
    doc.add_page_break()

    add_heading(doc, "Risk Matrix", 1)
    add_matrix(doc, ["Risk", "Likelihood", "Impact", "Mitigation"], [
        ["Domain registered under developer email", "Medium", "Critical", "Client-owned registrar; transfer lock; recovery records"],
        ["Live payment keys not switched", "Medium", "High", "Separate test/live env vars and launch-day payment test"],
        ["Database open to public IPs", "Low", "Critical", "Least-privilege users, IP controls, alerts, backups"],
        ["Images unoptimized", "Medium", "Medium", "Cloudinary transformations and responsive image testing"],
        ["SEO verification skipped", "Medium", "Medium", "Search Console, sitemap, robots, canonical and schema checks"],
        ["No maintenance owner", "High", "High", "Signed maintenance calendar and support escalation workflow"],
    ], [1.7, .9, .9, 3.0], header_fill=LIGHT_GRAY)
    doc.add_page_break()

    add_heading(doc, "Estimated Annual Running Cost", 1)
    add_matrix(doc, ["Service", "Low Budget", "Recommended", "Enterprise / High Growth"], [
        ["Domain", "INR 800-1,500", "INR 1,200-3,500", "INR 3,500+ with premium DNS"],
        ["Vercel Hosting", "Free tier", "Pro as traffic/team grows", "Team/Enterprise plan"],
        ["MongoDB Atlas", "Free M0", "Shared/dedicated cluster", "Dedicated cluster with advanced backups"],
        ["Cloudinary", "Free tier", "Paid media plan", "Advanced DAM and transformations"],
        ["Firebase Auth", "Usually free initially", "Usage-based", "Usage-based with monitoring"],
        ["Google Workspace", "Optional personal Gmail", "Business Starter/Standard", "Business Plus/Enterprise"],
        ["Razorpay", "Transaction fee", "Transaction fee", "Custom rate negotiation"],
        ["Maintenance", "INR 36,000/year", "INR 96,000-180,000/year", "INR 300,000+/year"],
        ["Estimated Total", "INR 40,000-70,000/year", "INR 120,000-250,000/year", "INR 400,000+/year"],
    ], [1.45, 1.55, 1.75, 1.75], header_fill=LIGHT_BLUE)
    doc.add_page_break()

    add_heading(doc, "Client Account Ownership Matrix", 1)
    add_matrix(doc, ["Service", "Purpose", "Who Should Own It", "Developer Access", "Client Access", "Criticality", "Renewal"], [
        ["Domain Registrar", "Domain purchase and DNS", "Client", "Admin or delegated DNS", "Owner", "Critical", "Yearly"],
        ["Vercel", "Hosting and deploys", "Client or agency-managed with agreement", "Project admin", "Owner/admin", "Critical", "Monthly/yearly"],
        ["MongoDB Atlas", "Database", "Client", "Project admin during support", "Org owner", "Critical", "Usage-based"],
        ["Cloudinary", "Media storage", "Client", "Admin during support", "Owner", "High", "Usage-based"],
        ["Firebase", "Authentication", "Client", "Editor/admin during support", "Owner", "High", "Usage-based"],
        ["Razorpay", "Payments", "Client", "Limited technical/webhook access", "Owner", "Critical", "Transaction-based"],
        ["Google Analytics", "Traffic insights", "Client", "Admin/editor", "Owner", "Medium", "Free"],
        ["Search Console", "Search indexing", "Client", "Full user", "Owner", "High", "Free"],
    ], [1.0, 1.1, 1.25, 1.1, 1.0, .8, .9], header_fill=LIGHT_GRAY)
    doc.add_page_break()

    add_heading(doc, "Complete Deployment Flowchart", 1)
    add_para(doc, "Mermaid source for documentation, README, or project wiki:")
    add_para(doc, """flowchart TD
    A[Business Decision] --> B[Create Client-Owned Accounts]
    B --> C[Purchase Domain]
    C --> D[Setup Vercel Hosting]
    D --> E[Setup MongoDB Atlas]
    E --> F[Configure Cloudinary]
    F --> G[Configure Firebase Authentication]
    G --> H[Connect Razorpay]
    H --> I[Configure Analytics and Search Console]
    I --> J[Deploy Production Build]
    J --> K[Full Testing]
    K --> L[SEO Verification]
    L --> M[Launch]
    M --> N[Maintenance Calendar]""")
    doc.add_page_break()

    add_heading(doc, "Maintenance Calendar", 1)
    add_matrix(doc, ["Frequency", "Activities", "Evidence to Record"], [
        ["Weekly", "Check forms, payments, Search Console issues, uptime, failed auth/payment events", "Maintenance log entry"],
        ["Monthly", "Review analytics, performance, broken links, backups, billing, dependency advisories", "Monthly health report"],
        ["Quarterly", "Update dependencies, review security, test restore plan, refresh SEO content, audit access", "Quarterly technical review"],
        ["Yearly", "Renew domain, review hosting plans, audit subscriptions, update business details", "Annual renewal checklist"],
    ], [1.1, 3.4, 2.0], header_fill=LIGHT_BLUE)
    doc.add_page_break()

    add_heading(doc, "Final Summary", 1)
    add_para(doc, "A professional jewellery website launch is not just a deploy button. It is a controlled business handover covering identity, ownership, infrastructure, data, media, authentication, payments, search, security, testing, support, and future growth.")
    add_para(doc, "The safest launch path is client-owned accounts, developer-assisted configuration, documented recovery, live-environment testing, and a maintenance agreement that begins immediately after launch.")


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_doc(doc)
    cover(doc)
    toc(doc)
    for num, title, topics, owner in PHASES:
        phase_content(doc, num, title, topics, owner)
    special_sections(doc)
    doc.save(OUT)


if __name__ == "__main__":
    build()
